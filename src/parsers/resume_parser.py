"""
src/parsers/resume_parser.py
-----------------------------
Parse resumes from PDF, DOCX, or plain text into a clean string.
"""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_resume(file_path: str | Path) -> str:
    """
    Extract plain text from a resume file.

    Supports: .pdf, .docx, .doc, .txt

    Returns:
        Extracted text as a single string.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use PDF, DOCX, or TXT.")


def parse_resume_bytes(content: bytes, filename: str) -> str:
    """Parse resume from raw bytes (for API file uploads)."""
    import tempfile, os
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        return parse_resume(tmp_path)
    finally:
        os.unlink(tmp_path)


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("PDF appears to be image-based (no selectable text). Use a text-based PDF.")
        logger.info("Parsed PDF: %d pages, %d chars", len(pages), len(text))
        return text
    except ImportError:
        raise ImportError("pypdf not installed. Run: pip install pypdf")


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables (some resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs).strip()
        logger.info("Parsed DOCX: %d paragraphs, %d chars", len(paragraphs), len(text))
        return text
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
