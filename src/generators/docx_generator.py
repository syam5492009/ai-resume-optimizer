"""
src/generators/docx_generator.py
----------------------------------
Generate an ATS-friendly DOCX from ResumeData.

ATS rules applied:
- No tables, text boxes, or frames
- No multi-column layout
- Standard heading styles (Word Heading 1/2)
- Simple fonts: Calibri for body, Calibri for headers
- 1-inch margins
- No images, no headers/footers
- Proper list styles for bullets
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.optimizer.rewriter import ResumeData


# ── Colour palette ────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x45, 0x6E)       # section headings
DARK = RGBColor(0x22, 0x22, 0x22)       # body text
MID  = RGBColor(0x55, 0x55, 0x55)       # secondary info


def generate_docx(resume: ResumeData, output_path: str | Path) -> Path:
    """
    Generate an ATS-optimised DOCX file.

    Args:
        resume:      Structured resume data.
        output_path: Where to save the .docx file.

    Returns:
        Path to the generated file.
    """
    doc = Document()
    _set_margins(doc)
    _set_default_font(doc)

    _add_header(doc, resume)
    _add_summary(doc, resume.summary)
    _add_skills(doc, resume.skills)
    _add_experience(doc, resume.experience)
    _add_projects(doc, resume.projects)
    if resume.pocs:
        _add_pocs(doc, resume.pocs)
    _add_education(doc, resume.education)
    if resume.certifications:
        _add_list_section(doc, "Certifications", resume.certifications)
    if resume.publications:
        _add_list_section(doc, "Publications", resume.publications)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ── Section builders ──────────────────────────────────────────────────────────

def _add_header(doc: Document, resume: ResumeData):
    """Name, title, and contact line — centred, no text box."""
    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(resume.name.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    _set_spacing(name_p, before=0, after=2)

    # Professional title
    if resume.title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(resume.title)
        run.font.size = Pt(11)
        run.font.color.rgb = MID
        run.bold = True
        _set_spacing(title_p, before=0, after=2)

    # Contact line
    c = resume.contact
    parts = [x for x in [c.location, c.phone, c.email] if x]
    links = [x for x in [c.portfolio, c.linkedin, c.github] if x]
    contact_line = " | ".join(parts)
    link_line = " | ".join(links)

    if contact_line:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(contact_line)
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK
        _set_spacing(cp, before=0, after=0)

    if link_line:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = lp.add_run(link_line)
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK
        _set_spacing(lp, before=0, after=4)

    _add_divider(doc)


def _add_summary(doc: Document, summary: str):
    if not summary:
        return
    _add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph(summary)
    p.style.font.size = Pt(10)
    _set_spacing(p, before=0, after=6)


def _add_skills(doc: Document, skills: dict[str, list[str]]):
    if not skills:
        return
    _add_section_heading(doc, "Core Technical Skills")
    for category, items in skills.items():
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=2)
        label = p.add_run(f"{category}: ")
        label.bold = True
        label.font.size = Pt(10)
        val = p.add_run(", ".join(items))
        val.font.size = Pt(10)


def _add_experience(doc: Document, experience: list):
    if not experience:
        return
    _add_section_heading(doc, "Professional Experience")
    for exp in experience:
        # Company + dates on same line
        p = doc.add_paragraph()
        _set_spacing(p, before=4, after=0)
        co = p.add_run(exp.company)
        co.bold = True
        co.font.size = Pt(11)
        co.font.color.rgb = NAVY
        if exp.start_date or exp.end_date:
            dates = f"{exp.start_date} – {exp.end_date}".strip(" –")
            dt = p.add_run(f"   |   {dates}")
            dt.font.size = Pt(10)
            dt.font.color.rgb = MID

        # Title + location
        tl = doc.add_paragraph()
        _set_spacing(tl, before=0, after=1)
        tr = tl.add_run(exp.title)
        tr.bold = True
        tr.font.size = Pt(10)
        if exp.location:
            lr = tl.add_run(f"  ·  {exp.location}")
            lr.font.size = Pt(10)
            lr.font.color.rgb = MID

        # Bullets
        for bullet in exp.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(bullet).font.size = Pt(10)
            _set_spacing(bp, before=0, after=1)


def _add_projects(doc: Document, projects: list):
    if not projects:
        return
    _add_section_heading(doc, "Technical Project Portfolio")
    for proj in projects:
        p = doc.add_paragraph()
        _set_spacing(p, before=4, after=0)
        nr = p.add_run(proj.name)
        nr.bold = True
        nr.font.size = Pt(10.5)
        nr.font.color.rgb = NAVY
        if proj.role:
            p.add_run(f"  |  {proj.role}").font.size = Pt(10)
        if proj.period:
            pr = p.add_run(f"  |  {proj.period}")
            pr.font.size = Pt(10)
            pr.font.color.rgb = MID

        if proj.url:
            up = doc.add_paragraph()
            _set_spacing(up, before=0, after=0)
            ur = up.add_run(f"Live: {proj.url}")
            ur.font.size = Pt(9.5)
            ur.font.color.rgb = MID
            ur.italic = True

        if proj.technologies:
            tp = doc.add_paragraph()
            _set_spacing(tp, before=0, after=0)
            tlabel = tp.add_run("Technologies: ")
            tlabel.bold = True
            tlabel.font.size = Pt(9.5)
            tp.add_run(proj.technologies).font.size = Pt(9.5)

        if proj.description:
            dp = doc.add_paragraph(proj.description)
            dp.runs[0].font.size = Pt(10)
            _set_spacing(dp, before=1, after=2)


def _add_pocs(doc: Document, pocs: list):
    _add_section_heading(doc, "Proof of Concepts (POCs) & R&D")
    for poc in pocs:
        p = doc.add_paragraph()
        _set_spacing(p, before=4, after=0)
        nr = p.add_run(poc.name)
        nr.bold = True
        nr.font.size = Pt(10.5)
        nr.font.color.rgb = NAVY
        if poc.description:
            dp = doc.add_paragraph(poc.description)
            dp.runs[0].font.size = Pt(10)
            _set_spacing(dp, before=1, after=2)


def _add_education(doc: Document, education: list):
    if not education:
        return
    _add_section_heading(doc, "Education")
    for edu in education:
        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=1)
        degree = f"{edu.degree}" + (f" in {edu.field}" if edu.field else "")
        dr = p.add_run(degree)
        dr.bold = True
        dr.font.size = Pt(10.5)
        if edu.institution:
            ir = p.add_run(f"  –  {edu.institution}")
            ir.font.size = Pt(10)
        if edu.year:
            yr = p.add_run(f"  |  {edu.year}")
            yr.font.size = Pt(10)
            yr.font.color.rgb = MID


def _add_list_section(doc: Document, title: str, items: list[str]):
    _add_section_heading(doc, title)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10)
        _set_spacing(p, before=0, after=1)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=8, after=2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    # Bottom border (thin rule under heading)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F456E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_divider(doc: Document):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F456E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(p, before: int = 0, after: int = 4):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def _set_default_font(doc: Document):
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = DARK
